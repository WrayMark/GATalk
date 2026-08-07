from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "GATalk_使用手册.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(90, 90, 90)
BODY_FONT = "Calibri"
CJK_FONT = "Microsoft YaHei UI"


def set_font(run, *, size: float, bold: bool = False, color=None) -> None:
    run.font.name = BODY_FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), BODY_FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), BODY_FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), CJK_FONT)
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def configure_styles(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = BODY_FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = styles[name]
        style.font.name = BODY_FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_numbering(document: Document, *, bullet: bool) -> int:
    numbering = document.part.numbering_part.element
    abstract_ids = [
        int(element.get(qn("w:abstractNumId")))
        for element in numbering.findall(qn("w:abstractNum"))
    ]
    abstract_id = max(abstract_ids, default=-1) + 1
    num_ids = [
        int(element.get(qn("w:numId")))
        for element in numbering.findall(qn("w:num"))
    ]
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_format = OxmlElement("w:numFmt")
    num_format.set(qn("w:val"), "bullet" if bullet else "decimal")
    level.append(num_format)
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "•" if bullet else "%1.")
    level.append(level_text)
    justification = OxmlElement("w:lvlJc")
    justification.set(qn("w:val"), "left")
    level.append(justification)
    paragraph_properties = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    paragraph_properties.append(tabs)
    indentation = OxmlElement("w:ind")
    indentation.set(qn("w:left"), "540")
    indentation.set(qn("w:hanging"), "270")
    paragraph_properties.append(indentation)
    level.append(paragraph_properties)
    abstract.append(level)
    numbering.append(abstract)

    number = OxmlElement("w:num")
    number.set(qn("w:numId"), str(num_id))
    abstract_reference = OxmlElement("w:abstractNumId")
    abstract_reference.set(qn("w:val"), str(abstract_id))
    number.append(abstract_reference)
    numbering.append(number)
    return num_id


def add_list_item(
    document: Document,
    text: str,
    num_id: int,
) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25
    properties = paragraph._p.get_or_add_pPr()
    numbering_properties = OxmlElement("w:numPr")
    level = OxmlElement("w:ilvl")
    level.set(qn("w:val"), "0")
    number = OxmlElement("w:numId")
    number.set(qn("w:val"), str(num_id))
    numbering_properties.append(level)
    numbering_properties.append(number)
    properties.append(numbering_properties)
    set_font(paragraph.add_run(text), size=11)


def add_body(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    set_font(paragraph.add_run(text), size=11)


def add_heading(document: Document, text: str, level: int = 1) -> None:
    paragraph = document.add_paragraph(text, style=f"Heading {level}")
    for run in paragraph.runs:
        set_font(
            run,
            size={1: 16, 2: 13, 3: 12}[level],
            bold=True,
            color=BLUE if level < 3 else DARK_BLUE,
        )


def build() -> None:
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    configure_styles(document)
    bullet_id = add_numbering(document, bullet=True)

    title = document.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(4)
    set_font(
        title.add_run("GATalk 简明使用手册"),
        size=22,
        bold=True,
        color=DARK_BLUE,
    )
    metadata = document.add_paragraph()
    metadata.paragraph_format.space_after = Pt(14)
    set_font(
        metadata.add_run(
            "适用版本：0.11.0（生产级资产拆分规划）"
            "　更新日期：2026-08-07"
        ),
        size=9.5,
        color=MUTED,
    )
    add_body(
        document,
        "本手册只说明当前可用功能、操作入口和必要限制。以后每次用户可见更新都会同步修订。",
    )

    add_heading(document, "1. 启动")
    add_list_item(
        document,
        "试用版：打开本次候选目录中的 GATalk/GATalk.exe，不要继续使用旧候选。",
        bullet_id,
    )
    add_list_item(document, "开发版：双击 start_dev.cmd。", bullet_id)
    add_list_item(
        document,
        "原始图片始终只读；GATalk 不会自动联网或上传图片。",
        bullet_id,
    )

    add_heading(document, "2. 工作台首页")
    for text in (
        "场景美术控制：参考图与自己的 UE 截图对比、审阅、任务和版本复查。",
        "作品研究：分析一张原画、概念图、Color Key 或优秀场景作品。",
        "资产拆分工作台：把复杂场景原画整理为资产清单、区域、生成图和展示板。",
        "模块内选择“文件 → 工作台首页”可返回。",
    ):
        add_list_item(document, text, bullet_id)

    add_heading(document, "3. 作品研究")
    add_heading(document, "最短流程", 2)
    decimal_id = add_numbering(document, bullet=False)
    for text in (
        "在首页选择“作品研究”。",
        "点击“新建作品研究”，选择目录并输入标题。",
        "导入一张图片。",
        "填写作品类型、本次想研究什么和已知背景。",
        "在“本地证据”查看明度、色板、空间网格和注意力代理。",
        "在“专家拆解”选择视觉 AI，查看发送清单后主动确认。",
        "阅读十二维拆解和跨维度因果链。",
        "在“学习笔记”写下自己的判断，保存。",
    ):
        add_list_item(document, text, decimal_id)
    add_heading(document, "本地观察", 2)
    for text in (
        "支持原图、灰度、三阶/五阶明度、伪色、溢出警告、剪影、缩略图、明度模糊和灯光明度代理图。",
        "构图辅助包括三分法、黄金分割、对角线、中心、三角形和透视线。",
        "点击 Oklab 色板颜色可查看来源区域；按 Esc 退出。",
        "注意力代理只综合局部反差、边缘密度和彩度，不是眼动、显著性或好坏判断。",
    ):
        add_list_item(document, text, bullet_id)
    add_heading(document, "AI 专家拆解", 2)
    for text in (
        "十二个维度：构图、视觉层级、明度、色彩、光、空间、形状、边缘细节、材质、环境叙事、风格技法和情绪。",
        "每项包含观察、画面证据、测量证据、解释、观看效果、评价取舍、学习点和不确定性。",
        "另外显示跨维度因果链、具体场景内容、画面标注、可迁移规律和继续观察问题。",
        "所有面向用户的 AI 内容固定使用简体中文，内部评价状态也会转换为中文。",
        "若首次返回英文或繁体中文，软件最多追加一次不含图片的中文规范化请求；发送确认窗口会提示可能增加费用。",
        "若 Gemini 当前模型重试后仍返回 503 容量不足，软件会按发送清单改用同一供应商的备用模型一次；完成状态会显示实际模型。",
        "旧版英文结果不会被本地硬翻译或继续展示，重新审阅即可生成中文结果。",
        "离线 Mock 只验证流程，不分析图片语义，也不代表本地 AI 推理。",
    ):
        add_list_item(document, text, bullet_id)
    add_heading(document, "保存内容", 2)
    add_body(
        document,
        ".scenelens-study 目录保存原图副本、SHA-256、本地证据、AI 结果、研究目标、已知背景、观察状态和个人笔记。原始图片不会被覆盖。",
    )

    add_heading(document, "4. 资产拆分工作台")
    add_heading(document, "从作品研究接续", 2)
    decimal_id = add_numbering(document, bullet=False)
    for text in (
        "在作品研究中完成本地证据或 AI 专家拆解后，点击“交给资产拆分…”。",
        "选择资产项目保存位置。GATalk 会复制原图原始字节并核对 SHA-256，同时带入研究目标、背景、笔记和分析摘要。",
        "进入资产拆分后先检查“作品研究交接”，可修改交接说明；本地文件路径不会发送给 AI。",
        "也可以直接新建资产项目并导入原画，不依赖作品研究。",
    ):
        add_list_item(document, text, decimal_id)
    add_heading(document, "先理解，再选择拆分程度", 2)
    decimal_id = add_numbering(document, bullet=False)
    for text in (
        "在“可校正拆分 → 场景理解与拆分方案”选择视觉供应商，查看发送清单并主动确认。",
        "先检查 AI 对空间层级、建筑或自然系统、重复规律、遮挡和不确定性的理解，并写入人工修正。",
        "从推荐方案中选择，或建立多个并存方案。常用预设包括空间总图、建筑／构筑物组、生产模块套件和细节构件。",
        "建筑、道具、植被、地形、材质贴花、远景等类别可分别设置拆分深度；深度 0 表示本方案不纳入该类别。",
        "确认方案后再生成资产清单。不同方案互不覆盖，可以针对同一原画分别得到整体建筑组、模块套件和门窗细件等结果。",
        "AI 建议只是起点；用户可修改方案名称、用途、每类深度、分组方式、单页上限和备注。",
    ):
        add_list_item(document, text, decimal_id)
    add_heading(document, "最短流程", 2)
    decimal_id = add_numbering(document, bullet=False)
    for text in (
        "在首页进入“资产拆分工作台”，新建资产项目。",
        "选择场景类型，填写制作目标，导入主原画；补充参考可选。",
        "先完成场景理解，选择并确认当前拆分方案。",
        "选择视觉供应商，查看资产清单发送内容并主动确认。",
        "检查资产分类、层级、复用组、优先级和原图区域。",
        "拖动区域或在“资产详情”修改名称、分类、父级、证据和制作策略。",
        "可新增、拆分、合并或删除资产；用户修订不会被后续 AI 覆盖。",
        "在清单第一列勾选需要生成的资产。",
        "在“生成与导出”选择独立概念图、保守遮挡补全图或评审展示图。",
        "查看发送清单并确认；取消时已完成项仍会保留。",
        "导出单项图片、多页 asset_board_01.png 等展示板和 asset_manifest.json。",
    ):
        add_list_item(document, text, decimal_id)
    add_heading(document, "画布与遮罩", 2)
    for text in (
        "“框选资产”后在原画拖出矩形；可移动或拉动手柄调整。",
        "“可见遮罩”只显示矩形内的算法近似像素，不是精确实例分割。",
        "按 Esc 退出框选或遮罩。原始图片不会被修改。",
    ):
        add_list_item(document, text, bullet_id)
    add_heading(document, "信息来源", 2)
    for text in (
        "原画可见证据：能直接从图中核对。",
        "AI 推断：类别、模块边界、复用或结构判断。",
        "用户补充／修订：用户确认后的内容，优先级最高。",
        "AI 生成补全：不可见部分的概念假设，不是原画事实。",
        "离线 Mock 只验证清单、保存、遮罩、生成和导出流程，不分析图片语义。",
    ):
        add_list_item(document, text, bullet_id)
    add_heading(document, "多方案与多页展示板", 2)
    for text in (
        "同一项目可以保存多个拆分方案。资产、生成记录和展示板都绑定所属方案。",
        "展示板可按资产家族／复用组、层级、空间系统或类别分组。内容超过单页上限时自动分页，不会强行缩小到一张难以阅读的图。",
        "AI 生成的独立资产图和展示板仍是概念辅助，不等于生产模型、真实背面或确定结构。",
    ):
        add_list_item(document, text, bullet_id)
    add_heading(document, "AI 清单结构修复", 2)
    for text in (
        "AI 偶尔会返回不存在的父资产 ID、重复 ID 或循环父级。新版会保留资产，只取消或修正无法成立的引用，不要求重新消耗一次 API 调用。",
        "修复发生时会显示中文摘要，并写入该次 AI 运行记录。",
        "GATalk 不会借此改写资产名称、分类、画面证据或制作建议。看到修复提示后，请在资产树中人工检查父子层级。",
    ):
        add_list_item(document, text, bullet_id)

    add_heading(document, "5. 场景美术控制最短流程")
    decimal_id = add_numbering(document, bullet=False)
    for text in (
        "点击“新建项目”，选择保存位置并命名。",
        "在左侧打开“制作意图”，填写目标风格、时间、天气、情绪、焦点和限制。",
        "新建 Shot，例如“村口固定机位”。",
        "导入参考图，再导入当前 UE 截图。后续截图会成为同 Shot 的新 Version。",
        "查看全图测量、共享色板和三阶明度差异。",
        "需要局部比较时，在“对比分析”中创建并配对矩形区域。",
        "将确认的问题转为修改任务。",
        "完成 UE 修改后导入新 Version，再进行复查。",
    ):
        add_list_item(document, text, decimal_id)
    add_body(document, "项目会自动保存。Ctrl+S 可立即保存。")

    add_heading(document, "6. 场景图片查看与基础分析")
    for text in (
        "滚轮缩放；左键拖动平移；双击画布恢复适配。",
        "开启“同步视图”可同步左右缩放和平移。",
        "选择“A/B 单图”后按 Space 切换参考图和当前截图。",
        "按 Esc 退出颜色遮罩、灯光标注、区域模式或优化预览。",
        "构图辅助提供三分法、黄金分割、对角线、中心、三角形、单点透视和两点透视；只用于人工观察。",
        "显示模式包括原图、灰度、三阶/五阶明度、曝光伪色、溢出警告、剪影、缩略图观察、明度模糊和灯光明度代理图。",
    ):
        add_list_item(document, text, bullet_id)
    add_body(
        document,
        "灯光明度代理图不是真正灰模，不能剥离材质和纹理。",
    )

    add_heading(document, "7. 对比分析")
    add_heading(document, "共享色板", 2)
    add_body(
        document,
        "使用同一组 Oklab 颜色比较双方占比。点击颜色可查看它在左右画面中的来源区域；再次点击或按 Esc 退出。",
    )
    add_heading(document, "三阶明度", 2)
    add_body(
        document,
        "显示暗部、中间调、亮部比例及百分点差。可调整阈值。这里只显示测量结果，不自动判断好坏。",
    )
    add_heading(document, "成对区域", 2)
    decimal_id = add_numbering(document, bullet=False)
    for text in (
        "进入“区域模式”。",
        "分别在参考图和当前截图拖出矩形。",
        "选中两侧区域，建立配对并设置名称和语义。",
        "松开鼠标后自动分析。",
    ):
        add_list_item(document, text, decimal_id)
    add_body(
        document,
        "区域可移动、缩放、删除和保存。新 Version 可复制上一版本区域，复制后必须人工检查位置。",
    )

    add_heading(document, "8. 制作意图与参考图视觉简报")
    for text in (
        "“制作意图”记录你希望最终画面达到什么目标。",
        "“参考图视觉简报”记录参考图实际呈现的视觉特征。",
        "自动测量、算法推断、AI 分析和用户填写会标记不同来源。",
        "用户确认或修订过的内容不会被算法或 AI 自动覆盖。",
    ):
        add_list_item(document, text, bullet_id)

    add_heading(document, "9. AI 审阅与灯光审片")
    decimal_id = add_numbering(document, bullet=False)
    for text in (
        "优先选择“深度主美审阅（八维）”；专项灯光问题选择“灯光专项审阅”。",
        "无 API Key 时选择“离线 Mock”验证流程。",
        "使用真实供应商前，把 Key 存入 Windows 系统凭据。",
        "点击“查看发送清单并审阅”，核对数据后手动确认发送。",
        "查看核心问题、证据支持或冲突，再由用户确认转为任务。",
    ):
        add_list_item(document, text, decimal_id)
    add_body(document, "“第二意见”会增加一次模型调用和费用，默认关闭。")
    add_heading(document, "深度主美审阅", 2)
    for text in (
        "八个维度：构图、视觉引导、焦点层级、色彩、明度、灯光、材质和世界设计。",
        "每个维度分开显示制作目标、参考呈现、当前效果、优点、风险和不确定性。",
        "最多五个核心问题，并给出保留项、UE 执行顺序和下一版验证方法。",
        "本地证据显示测量支持、部分支持、存在冲突或无法验证；冲突不会被隐藏。",
        "离线 Mock 只验证界面和 JSON 流程，不会在本地分析图片，也不代表真实审阅。",
    ):
        add_list_item(document, text, bullet_id)
    add_heading(document, "AI 接口失败时", 2)
    for text in (
        "connection_closed：远端或中间网络设备提前断开连接，不表示 API Key 错误。新版会自动退避重试 3 次；仍失败时检查代理、VPN 和网络稳定性。",
        "自动重试可能重复提交同一请求并产生额外费用，发送确认窗口会提前提示。",
        "http_400：请求参数或模型不兼容。新版会自动处理 Gemini 深度审阅的复杂 Schema；仍失败时先确认正在运行 0.5.0a4，再恢复默认模型重试。",
        "若错误包含 generation_config.response_format.text.mime_type，说明仍在运行旧版；关闭旧程序后改用本次 Gemini 接口修正版。",
        "Gemini 首次返回的 JSON 语法损坏、被截断或结构不完整时，GATalk 最多自动纠错一次。确认发送窗口会提示该过程可能再次发送同一审阅副本并增加少量费用。",
        "finish_reason=MAX_TOKENS：AI 回答过长导致 JSON 尾部被截断，不表示 API Key 错误。新版会压缩内容，并只在明确截断时用更高预算纠错一次。",
        "错误中的 line 和 column 表示 AI 返回文本的 JSON 损坏位置，不表示 API Key 错误。0.5.0a4 会把原始返回交给同一模型压缩并修复；第二次仍损坏时停止，不会无限重试。",
        "若 AI 引用了本次报告中不存在的问题 ID，0.5.0a3 会取消无效链接并继续显示完整报告，不追加 API 调用。界面的“结构修复”只表示链接已取消，正文和真实问题没有被改写。",
        "http_401：Key 错误、过期或不属于该供应商。",
        "http_403：检查模型权限、服务开通、地域和 Key 所属工作区。",
        "http_404：模型 ID 已下线或接口不存在。Gemini 会按发送清单自动尝试当前备用链；其他供应商请改回其默认模型。",
        "http_413：降低发送分辨率后重试。",
        "http_429：额度不足或调用过快，检查控制台后稍后重试。",
        "http_503：供应商当前模型临时过载，不表示 API Key 错误。新版先有限重试；Gemini 仍失败时会按发送清单依次尝试 gemini-3.5-flash 和 gemini-3.5-flash-lite，不再使用 gemini-2.5-flash。",
        "发送清单会提前显示完整备用链。若发生回退，状态栏和保存记录会显示请求模型、实际模型及回退原因；每次实际发送都可能产生调用费用。",
        "其他 http_5xx：供应商服务异常，稍后重试。",
    ):
        add_list_item(document, text, bullet_id)
    add_body(
        document,
        "错误窗口会显示供应商返回的脱敏原因。不要把 API Key 截图或复制给他人。",
    )

    add_heading(document, "10. 优化实验室")
    add_heading(document, "目标匹配画像", 2)
    add_body(
        document,
        "显示明度、黑白灰、色板、彩度、冷暖和区域关系等维度。可修改权重。“估计匹配度”只代表当前算法和权重，不是作品质量评分。证据不足的维度不会自动猜测。",
    )
    add_heading(document, "安全调色", 2)
    decimal_id = add_numbering(document, bullet=False)
    for text in (
        "调整曝光、对比、白平衡、阴影、中间调、亮部和彩度。",
        "选择全图或当前选中的配对区域。",
        "设置参考影响强度并生成预览。",
        "使用 A/B、撤销和重做检查结果。",
        "可导出 PNG 和 JSON 配方；符合条件时可导出 .cube。",
    ):
        add_list_item(document, text, decimal_id)
    add_body(
        document,
        "安全调色完全本地运行，不修改原图。区域调色或参考色迁移不能导出通用 .cube。",
    )
    add_heading(document, "AI 优化预演", 2)
    decimal_id = add_numbering(document, bullet=False)
    for text in (
        "选择只改灯光、只改色彩或只改雾与氛围。",
        "设置改动预算和构图、几何、资产身份保护项。",
        "查看发送清单并主动确认。",
        "检查结构漂移、构图偏移和保护区变化。",
    ):
        add_list_item(document, text, decimal_id)
    add_body(
        document,
        "AI 输出只保存为 AIConceptPreview，不会成为真实 UE Version。出现“仅适合概念参考”时，不应把预演当作可直接复现的场景结果。",
    )

    add_heading(document, "11. 文件与安全")
    for text in (
        "assets：导入原图，保留原始字节和 SHA-256。",
        "artifacts：可重建分析结果和 AIConceptPreview。",
        "exports：审阅包、预览和配方。",
        "API Key 不写入项目、SQLite、JSON、日志或 Git。",
        "同一项目只能由一个进程写入；第二个进程可只读打开。",
        "作品研究的 assets 同样保留导入图片原始字节；study.json 保存研究状态。",
        "资产拆分项目使用 .scenelens-assets 目录；生成图、遮罩和导出记录与原图分开保存。",
    ):
        add_list_item(document, text, bullet_id)

    add_heading(document, "12. 当前限制")
    for text in (
        "只支持静态图片和矩形区域。",
        "不支持视频、HDR/EXR、多边形、自动分割或 UE 工程扫描。",
        "真实 AI 供应商需要单独验证账号、地区、模型和费用。",
        "AI 预演必须回到 UE 实施，并用新的真实截图 Version 正式复查。",
        "作品研究暂不支持多图研究、引用资料库、自动语义分割或本地大型视觉模型。",
        "资产拆分不捆绑本地视觉大模型、SAM 2 或 CUDA，也不生成生产可用三维模型。",
        "AI 资产分类、遮挡补全和生成图仍需用户人工校正。",
    ):
        add_list_item(document, text, bullet_id)

    document.core_properties.title = "GATalk 简明使用手册"
    document.core_properties.subject = "GATalk 当前功能与操作"
    document.core_properties.author = "GATalk"
    document.core_properties.keywords = "GATalk, 使用手册, 游戏场景美术"
    document.save(OUTPUT)


if __name__ == "__main__":
    build()
